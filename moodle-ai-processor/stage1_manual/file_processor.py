"""
Stage 1: Manual File Processing

This module processes uploaded HTML and Word documents from Moodle
using AI without direct API integration.
"""

import os
import json
import logging
import time
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import re

from bs4 import BeautifulSoup
import docx
import sys
sys.path.append(str(Path(__file__).parent.parent / 'src'))
from src.ai_client import OpenRouterClient


class FileProcessor:
    """Process uploaded HTML and Word documents"""
    
    def __init__(self, config_path: str = "config/config.json"):
        """Initialize the file processor"""
        # Handle both relative and absolute config paths
        if not Path(config_path).exists():
            config_path = Path(__file__).parent.parent / config_path
        
        self.config_path = config_path
        self.config = self._load_config()
        self.ai_client = OpenRouterClient(
            api_key=self.config['openrouter']['api_key'],
            base_url=self.config['openrouter']['base_url'],
            model=self.config['openrouter']['model']
        )
        
        # Set paths relative to the project root
        project_root = Path(__file__).parent.parent
        self.upload_dir = project_root / "stage1_manual/uploads"
        self.processed_dir = project_root / "stage1_manual/processed"
        self.logger = logging.getLogger(__name__)
        
        # Ensure directories exist
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
    
    def _load_config(self) -> Dict[str, Any]:
        """Load configuration from JSON file"""
        with open(self.config_path, 'r') as f:
            return json.load(f)
    
    def extract_text_from_html(self, html_content: str) -> Dict[str, Any]:
        """Extract text and structure from HTML"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract forum posts if it's a forum page
        posts = []
        
        # Common Moodle forum post selectors
        post_elements = soup.find_all(['div', 'article'], class_=re.compile(r'post|discussion|forum'))
        
        if not post_elements:
            # Fallback: look for any content that might be posts
            post_elements = soup.find_all(['div', 'p'], string=re.compile(r'.{50,}'))
        
        for element in post_elements:
            post_text = element.get_text().strip()
            if len(post_text) > 50:  # Filter out very short content
                # Try to extract author and timestamp
                author = self._extract_author(element)
                timestamp = self._extract_timestamp(element)
                
                posts.append({
                    'content': post_text,
                    'author': author,
                    'timestamp': timestamp,
                    'html_class': element.get('class', [])
                })
        
        # If no posts found, extract all meaningful text
        if not posts:
            text = soup.get_text()
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            posts.append({
                'content': text,
                'author': 'Unknown',
                'timestamp': None,
                'html_class': []
            })
        
        return {
            'posts': posts,
            'total_posts': len(posts),
            'extraction_method': 'html_parsing'
        }
    
    def _extract_author(self, element) -> str:
        """Try to extract author from HTML element"""
        # Look for common author indicators
        author_element = element.find(['span', 'div', 'p'], class_=re.compile(r'author|user|name'))
        if author_element:
            return author_element.get_text().strip()
        
        # Look for "by" patterns
        text = element.get_text()
        by_match = re.search(r'by\s+([^,\n]+)', text, re.IGNORECASE)
        if by_match:
            return by_match.group(1).strip()
        
        return 'Unknown'
    
    def _extract_timestamp(self, element) -> Optional[str]:
        """Try to extract timestamp from HTML element"""
        # Look for time elements
        time_element = element.find('time')
        if time_element:
            return time_element.get('datetime') or time_element.get_text().strip()
        
        # Look for date patterns
        text = element.get_text()
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{4}-\d{2}-\d{2}',
            r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        
        return None
    
    def extract_posts_from_json(self, json_content: str) -> Dict[str, Any]:
        """Extract posts from JSON discussion data, filtering personal information"""
        try:
            # Parse JSON content
            if json_content.strip().startswith('[['):
                # Handle nested array structure
                data = json.loads(json_content)[0]  # Get first array element
            else:
                data = json.loads(json_content)
            
            posts = []
            personal_info_fields = ['userid', 'userfullname', 'privatereplytofullname']
            
            for post_data in data:
                # Extract safe content (no personal info)
                safe_post = {
                    'id': post_data.get('id'),
                    'discussion': post_data.get('discussion'),
                    'parent': post_data.get('parent', 0),
                    'created': post_data.get('created'),
                    'modified': post_data.get('modified'),
                    'subject': post_data.get('subject', ''),
                    'message': post_data.get('message', ''),
                    'wordcount': post_data.get('wordcount', 0),
                    'charcount': post_data.get('charcount', 0)
                }
                
                # Clean HTML from message content
                if safe_post['message']:
                    soup = BeautifulSoup(safe_post['message'], 'html.parser')
                    safe_post['clean_message'] = soup.get_text().strip()
                else:
                    safe_post['clean_message'] = ''
                
                # Convert timestamp if needed
                if safe_post['created']:
                    try:
                        safe_post['created_readable'] = datetime.fromtimestamp(safe_post['created']).isoformat()
                    except (ValueError, OSError):
                        safe_post['created_readable'] = safe_post['created']
                
                # Only include posts with meaningful content
                if safe_post['clean_message'] and len(safe_post['clean_message']) > 10:
                    posts.append(safe_post)
            
            print(f"   🔒 Filtered out personal information from {len(data)} posts")
            print(f"   ✅ Extracted {len(posts)} posts with content")
            
            return {
                'posts': posts,
                'total_posts': len(posts),
                'extraction_method': 'json_parsing',
                'personal_info_filtered': True
            }
            
        except json.JSONDecodeError as e:
            self.logger.error(f"Error parsing JSON: {e}")
            return {'error': f'JSON parsing error: {str(e)}'}
        except Exception as e:
            self.logger.error(f"Error extracting from JSON: {e}")
            return {'error': str(e)}
    
    def extract_text_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = docx.Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'content': para.text.strip(),
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Also extract from tables
            tables_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        tables_content.append(row_data)
            
            return {
                'paragraphs': paragraphs,
                'tables': tables_content,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables_content),
                'extraction_method': 'docx_parsing'
            }
        except Exception as e:
            self.logger.error(f"Error extracting from docx: {e}")
            return {'error': str(e)}
    
    def process_file(self, file_path: Path, processing_type: str = 'feedback') -> Dict[str, Any]:
        """Process a single file with AI"""
        print(f"\n📄 Processing file: {file_path.name}")
        print(f"   Type: {processing_type}")
        print(f"   Size: {file_path.stat().st_size:,} bytes")
        
        result = {
            'file_name': file_path.name,
            'file_path': str(file_path),
            'processed_at': datetime.now().isoformat(),
            'processing_type': processing_type
        }
        
        try:
            # Extract content based on file type
            if file_path.suffix.lower() == '.json':
                print("   🔍 Extracting content from JSON...")
                with open(file_path, 'r', encoding='utf-8') as f:
                    json_content = f.read()
                extracted_data = self.extract_posts_from_json(json_content)
                result['extraction'] = extracted_data
                
                if 'error' in extracted_data:
                    result['error'] = extracted_data['error']
                    return result
                
                print(f"   ✅ Found {extracted_data['total_posts']} posts")
                
                # Process each post
                processed_posts = []
                for i, post in enumerate(extracted_data['posts'], 1):
                    print(f"\n   📝 Processing post {i}/{len(extracted_data['posts'])}")
                    print(f"      Subject: {post['subject'][:50]}...")
                    print(f"      Length: {len(post['clean_message'])} chars")
                    
                    ai_response = self._process_content_with_ai(
                        content=post['clean_message'],
                        author=f"User_{post['id']}",  # Anonymized reference
                        processing_type=processing_type
                    )
                    processed_posts.append({
                        'original': post,
                        'ai_response': ai_response
                    })
                    print(f"   ✅ Post {i} processed successfully")
                
                result['processed_posts'] = processed_posts
                print(f"\n✅ JSON processing complete: {len(processed_posts)} posts processed")
                
            elif file_path.suffix.lower() == '.html':
                print("   🔍 Extracting content from HTML...")
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                extracted_data = self.extract_text_from_html(html_content)
                result['extraction'] = extracted_data
                print(f"   ✅ Found {extracted_data['total_posts']} posts/sections")
                
                # Process each post
                processed_posts = []
                for i, post in enumerate(extracted_data['posts'], 1):
                    print(f"\n   📝 Processing post {i}/{len(extracted_data['posts'])}")
                    print(f"      Author: {post['author']}")
                    print(f"      Length: {len(post['content'])} chars")
                    
                    ai_response = self._process_content_with_ai(
                        content=post['content'],
                        author=post['author'],
                        processing_type=processing_type
                    )
                    processed_posts.append({
                        'original': post,
                        'ai_response': ai_response
                    })
                    print(f"   ✅ Post {i} processed successfully")
                
                result['processed_posts'] = processed_posts
                print(f"\n✅ HTML processing complete: {len(processed_posts)} posts processed")
                
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                print("   🔍 Extracting content from Word document...")
                extracted_data = self.extract_text_from_docx(file_path)
                result['extraction'] = extracted_data
                
                if 'error' not in extracted_data:
                    print(f"   ✅ Found {extracted_data['total_paragraphs']} paragraphs")
                    # Combine all paragraphs for AI processing
                    full_content = '\n\n'.join([p['content'] for p in extracted_data['paragraphs']])
                    
                    print(f"   📝 Processing document content ({len(full_content)} chars)...")
                    ai_response = self._process_content_with_ai(
                        content=full_content,
                        processing_type=processing_type
                    )
                    result['ai_response'] = ai_response
                    print("   ✅ Document processing complete")
                else:
                    print(f"   ❌ Error extracting content: {extracted_data['error']}")
            
            else:
                print("   🔍 Processing as plain text...")
                # Try to read as plain text
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                print(f"   📝 Processing text content ({len(content)} chars)...")
                ai_response = self._process_content_with_ai(
                    content=content,
                    processing_type=processing_type
                )
                result['ai_response'] = ai_response
                print("   ✅ Text processing complete")
        
        except Exception as e:
            result['error'] = str(e)
            self.logger.error(f"Error processing {file_path}: {e}")
        
        return result
    
    def _process_content_with_ai(self, content: str, author: Optional[str] = None, 
                                processing_type: str = 'feedback') -> Dict[str, Any]:
        """Process content with AI based on processing type"""
        try:
            print(f"\n🤖 Starting AI processing...")
            print(f"   📋 Processing type: {processing_type}")
            if author:
                print(f"   👤 Author: {author}")
            print(f"   📝 Content length: {len(content)} characters")
            print(f"   📊 Content preview: {content[:100]}..." if len(content) > 100 else f"   📊 Content: {content}")
            
            if processing_type == 'feedback':
                print("\n   🔍 Step 1/2: Analyzing student post...")
                print("   🚀 Sending analysis request to AI...")
                analysis_start = time.time()
                analysis = self.ai_client.analyze_student_post(content)
                analysis_time = time.time() - analysis_start
                print(f"   ✅ Analysis complete in {analysis_time:.2f} seconds")
                
                print("\n   💬 Step 2/2: Generating reply...")
                print("   🚀 Sending reply generation request to AI...")
                reply_start = time.time()
                reply = self.ai_client.generate_reply(content, reply_style="supportive")
                reply_time = time.time() - reply_start
                print(f"   ✅ Reply generated in {reply_time:.2f} seconds")
                
                total_time = analysis_time + reply_time
                print(f"\n   🎯 Total AI processing time: {total_time:.2f} seconds")
                
                return {
                    'type': 'feedback',
                    'analysis': analysis,
                    'suggested_reply': reply,
                    'author': author,
                    'processing_time': total_time
                }
            
            elif processing_type == 'summary':
                print("\n   📋 Generating summary...")
                summary_prompt = f"Please provide a concise summary of this student work:\n\n{content}"
                print("   🚀 Sending summary request to AI...")
                summary_start = time.time()
                summary = self.ai_client.generate_response(
                    summary_prompt,
                    system_prompt="You are an educational assistant. Provide clear, concise summaries."
                )
                summary_time = time.time() - summary_start
                print(f"   ✅ Summary complete in {summary_time:.2f} seconds")
                
                return {
                    'type': 'summary',
                    'summary': summary,
                    'author': author,
                    'processing_time': summary_time
                }
            
            elif processing_type == 'grading':
                print("   📊 Generating grading feedback...")
                grading_prompt = f"""Please provide grading feedback for this student work:

{content}

Provide:
1. Strengths of the work
2. Areas for improvement  
3. Suggested grade/score (explain your reasoning)
4. Specific recommendations"""
                
                grading = self.ai_client.generate_response(
                    grading_prompt,
                    system_prompt="You are an experienced educator providing constructive grading feedback."
                )
                print("   ✅ Grading feedback complete")
                
                return {
                    'type': 'grading',
                    'grading_feedback': grading,
                    'author': author
                }
            
        except Exception as e:
            return {
                'type': processing_type,
                'error': str(e)
            }
    
    def process_upload_directory(self, processing_type: str = 'feedback') -> Dict[str, Any]:
        """Process all files in the upload directory"""
        upload_path = Path(self.upload_dir)
        results = []
        
        print(f"\n🚀 Starting batch processing (type: {processing_type})")
        print(f"📁 Upload directory: {upload_path}")
        
        # Get all supported files
        supported_extensions = ['.html', '.htm', '.docx', '.doc', '.txt', '.json']
        files_to_process = []
        
        for ext in supported_extensions:
            files_to_process.extend(upload_path.glob(f'*{ext}'))
        
        if not files_to_process:
            print("❌ No supported files found in upload directory")
            print(f"   Supported formats: {', '.join(supported_extensions)}")
            return {'error': 'No files to process'}
        
        print(f"📋 Found {len(files_to_process)} files to process")
        for i, file_path in enumerate(files_to_process, 1):
            print(f"   {i}. {file_path.name} ({file_path.stat().st_size:,} bytes)")
        
        print(f"\n⚡ Processing files...")
        
        for i, file_path in enumerate(files_to_process, 1):
            print(f"\n{'='*50}")
            print(f"📄 Processing file {i}/{len(files_to_process)}")
            print(f"{'='*50}")
            
            result = self.process_file(file_path, processing_type)
            results.append(result)
            
            # Save individual result
            result_file = self.processed_dir / f"{file_path.stem}_{processing_type}_result.json"
            print(f"💾 Saving results to: {result_file.name}")
            with open(result_file, 'w') as f:
                json.dump(result, f, indent=2, default=str)
            
            if 'error' in result:
                print(f"❌ Error processing {file_path.name}: {result['error']}")
            else:
                print(f"✅ Successfully processed {file_path.name}")
        
        # Save combined results
        print(f"\n📊 Generating combined results...")
        combined_result = {
            'processing_type': processing_type,
            'processed_at': datetime.now().isoformat(),
            'total_files': len(files_to_process),
            'results': results
        }
        
        combined_file = self.processed_dir / f"combined_{processing_type}_results.json"
        print(f"💾 Saving combined results to: {combined_file.name}")
        with open(combined_file, 'w') as f:
            json.dump(combined_result, f, indent=2, default=str)
        
        # Final summary
        success_count = len([r for r in results if 'error' not in r])
        error_count = len(files_to_process) - success_count
        
        print(f"\n🎉 Batch processing complete!")
        print(f"   ✅ Successfully processed: {success_count}")
        if error_count > 0:
            print(f"   ❌ Errors: {error_count}")
        print(f"   📁 Results saved to: {self.processed_dir}")
        
        return combined_result
    
    def export_markdown_feedback(self, results_data: Dict[str, Any]) -> str:
        """
        Export feedback as well-structured markdown with clearly divided sections
        that can be delivered together or separately
        """
        processing_type = results_data.get('processing_type', 'feedback')
        processed_at = results_data.get('processed_at', datetime.now().isoformat())
        
        markdown_content = f"""# 📝 AI-Generated Feedback Report

**Processing Type:** {processing_type.title()}  
**Generated:** {processed_at}  
**Total Files Processed:** {results_data.get('total_files', 0)}

---

"""
        
        # Process each file's results
        for file_idx, result in enumerate(results_data.get('results', []), 1):
            file_name = result.get('file_name', f'File {file_idx}')
            markdown_content += f"""## 📄 File {file_idx}: {file_name}

"""
            
            # Handle different file types and their processed posts
            if 'processed_posts' in result:
                # For HTML/JSON files with multiple posts
                posts = result['processed_posts']
                markdown_content += f"**Posts in this file:** {len(posts)}\n\n"
                
                for post_idx, post_data in enumerate(posts, 1):
                    original = post_data.get('original', {})
                    ai_response = post_data.get('ai_response', {})
                    
                    # Extract post identification
                    if 'subject' in original:
                        post_title = original.get('subject', f'Post {post_idx}')
                    else:
                        post_title = f"Post {post_idx}"
                    
                    # Create clearly divided section for each student
                    markdown_content += f"""### 🎓 {post_title}

<details>
<summary><strong>📖 Original Post</strong> (Click to expand)</summary>

```
{original.get('clean_message', original.get('content', 'No content available'))}
```

</details>

#### 💬 AI Feedback

"""
                    
                    if ai_response.get('type') == 'feedback':
                        # Format feedback response
                        analysis = ai_response.get('analysis', 'No analysis available')
                        reply = ai_response.get('suggested_reply', 'No reply available')
                        
                        markdown_content += f"""**Analysis:**
{analysis}

**Suggested Reply:**
{reply}

"""
                    else:
                        # Handle other response types
                        for key, value in ai_response.items():
                            if key not in ['type', 'author', 'processing_time']:
                                markdown_content += f"**{key.replace('_', ' ').title()}:**\n{value}\n\n"
                    
                    # Add separator between posts
                    markdown_content += "---\n\n"
            
            elif 'ai_response' in result:
                # For single document processing (like Word docs)
                ai_response = result['ai_response']
                markdown_content += f"""### 📝 Document Feedback

"""
                if ai_response.get('type') == 'feedback':
                    analysis = ai_response.get('analysis', 'No analysis available')
                    reply = ai_response.get('suggested_reply', 'No reply available')
                    
                    markdown_content += f"""**Analysis:**
{analysis}

**Suggested Response:**
{reply}

"""
                else:
                    for key, value in ai_response.items():
                        if key not in ['type', 'author', 'processing_time']:
                            markdown_content += f"**{key.replace('_', ' ').title()}:**\n{value}\n\n"
                
                markdown_content += "---\n\n"
        
        # Add footer
        markdown_content += f"""---

## 📋 Export Information

This feedback report was generated using AI assistance to support student learning and engagement. 

**How to use this report:**

1. **📤 Complete Report**: Copy the entire document for comprehensive feedback
2. **🎯 Individual Sections**: Copy specific student sections for personalized delivery
3. **📋 Consolidated Post**: Combine multiple sections for a forum summary post
4. **📧 Email Templates**: Use individual sections as email templates

**Delivery Options:**
- Copy sections to create individual Moodle forum replies
- Use as basis for email feedback to students
- Create a single comprehensive forum post
- Print for offline review and annotation

---
*Generated by Moodle AI Processor - Stage 1*
"""
        
        # Save markdown file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        markdown_file = self.processed_dir / f"feedback_report_{timestamp}.md"
        
        with open(markdown_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"📄 Markdown report saved: {markdown_file.name}")
        return str(markdown_file)
    
    def export_consolidated_forum_post(self, results_data: Dict[str, Any]) -> str:
        """
        Generate a single consolidated forum post with all feedback
        Ready for copy-paste into Moodle forum
        """
        processing_type = results_data.get('processing_type', 'feedback')
        
        post_content = f"""<h2>🤖 AI Teaching Assistant Feedback</h2>

<p><strong>Hello everyone!</strong></p>

<p>I've reviewed the recent posts and discussions, and I'd like to provide some constructive feedback and insights to support your learning journey.</p>

<hr>

"""
        
        student_count = 0
        for file_idx, result in enumerate(results_data.get('results', []), 1):
            if 'processed_posts' in result:
                posts = result['processed_posts']
                
                for post_idx, post_data in enumerate(posts, 1):
                    student_count += 1
                    original = post_data.get('original', {})
                    ai_response = post_data.get('ai_response', {})
                    
                    post_title = original.get('subject', f'Discussion Post {student_count}')
                    
                    post_content += f"""<h3>💬 Regarding: "{post_title}"</h3>

<div style="background-color: #f8f9fa; padding: 15px; border-left: 4px solid #007bff; margin: 15px 0;">
"""
                    
                    if ai_response.get('type') == 'feedback':
                        analysis = ai_response.get('analysis', '')
                        reply = ai_response.get('suggested_reply', '')
                        
                        post_content += f"""<p><strong>Analysis:</strong> {analysis}</p>
<p><strong>Response:</strong> {reply}</p>
"""
                    else:
                        for key, value in ai_response.items():
                            if key not in ['type', 'author', 'processing_time']:
                                post_content += f"<p><strong>{key.replace('_', ' ').title()}:</strong> {value}</p>"
                    
                    post_content += "</div>\n\n"
        
        post_content += f"""<hr>

<p><strong>Summary:</strong> This feedback covers {student_count} posts and is designed to help enhance our discussion and learning. Please feel free to respond with questions or continue the conversation!</p>

<p><em>This feedback was generated with AI assistance to support your learning and engagement.</em></p>
"""
        
        # Save HTML file for forum posting
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        html_file = self.processed_dir / f"forum_post_{timestamp}.html"
        
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(post_content)
        
        print(f"📤 Forum post ready: {html_file.name}")
        print(f"💡 Copy the contents and paste into your Moodle forum")
        
        return str(html_file)
